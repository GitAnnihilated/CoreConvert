"""
CoreConvert — Conversion library
"""
import os
import io
import subprocess

# ── Custom exception ──────────────────────────────────────────────────────────

class ConversionError(ValueError):
    pass


# ── Supported conversion map ──────────────────────────────────────────────────

SUPPORTED_TARGETS: dict[str, list[str]] = {
    # ── Data ──────────────────────────────────────────────────────────────────
    'csv':  ['json', 'xlsx'],
    'json': ['csv', 'xlsx'],
    'xml':  ['csv', 'json', 'xlsx'],
    'xlsx': ['csv', 'json'],
    'xls':  ['csv', 'json'],
    # ── Documents ─────────────────────────────────────────────────────────────
    'pdf':  ['docx', 'txt'],
    'docx': ['pdf', 'txt', 'html'],
    'html': ['pdf', 'md'],
    'md':   ['html'],
    # ── Images (raster) ───────────────────────────────────────────────────────
    'png':  ['jpg', 'gif', 'bmp', 'webp', 'pdf'],
    'jpg':  ['png', 'gif', 'bmp', 'webp', 'pdf'],
    'jpeg': ['png', 'gif', 'bmp', 'webp', 'pdf'],
    'gif':  ['png', 'jpg', 'bmp', 'webp'],
    'bmp':  ['png', 'jpg', 'gif', 'webp', 'pdf'],
    'webp': ['png', 'jpg', 'gif', 'bmp', 'pdf'],
    # ── Images (special) ──────────────────────────────────────────────────────
    'heic': ['jpg', 'png'],
    'svg':  ['png', 'pdf'],
    # ── Audio ─────────────────────────────────────────────────────────────────
    'mp3':  ['wav', 'ogg', 'flac'],
    'wav':  ['mp3', 'ogg', 'flac'],
    'ogg':  ['mp3', 'wav', 'flac'],
    'flac': ['mp3', 'wav', 'ogg'],
    'aac':  ['mp3', 'wav'],
    'm4a':  ['mp3', 'wav'],
    # ── Video ─────────────────────────────────────────────────────────────────
    'mp4':  ['gif', 'mp3', 'avi', 'mov', 'mkv'],
    'avi':  ['gif', 'mp3', 'mp4', 'mov', 'mkv'],
    'mov':  ['gif', 'mp3', 'mp4', 'avi', 'mkv'],
    'mkv':  ['gif', 'mp3', 'mp4', 'avi', 'mov'],
    'webm': ['gif', 'mp3', 'mp4'],
    # ── Archives ──────────────────────────────────────────────────────────────
    'zip':  ['tar'],
    'rar':  ['zip'],
    'tar':  ['zip'],
}

RASTER_IMAGE_FORMATS = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp'}
DATA_FORMATS         = {'json', 'xml', 'csv', 'xlsx', 'xls'}
VIDEO_FORMATS        = {'mp4', 'avi', 'mov', 'mkv', 'webm'}
AUDIO_FORMATS        = {'mp3', 'wav', 'ogg', 'flac', 'aac', 'm4a'}
ARCHIVE_FORMATS      = {'zip', 'rar', 'tar'}


# ── Router ────────────────────────────────────────────────────────────────────

def convert_file(input_path: str, output_path: str, target: str) -> None:
    src_ext = os.path.splitext(input_path)[1].lower().lstrip('.')

    allowed = SUPPORTED_TARGETS.get(src_ext)
    if allowed is None:
        raise ConversionError(
            f"Unsupported file type: .{src_ext}. "
            f"Supported source formats: {', '.join(sorted(SUPPORTED_TARGETS.keys()))}"
        )
    if target not in allowed:
        targets_str = ', '.join(f'.{t.upper()}' for t in allowed)
        raise ConversionError(
            f".{src_ext.upper()} files can be converted to: {targets_str}"
        )

    # ── Data ──────────────────────────────────────────────────────────────────
    if src_ext in DATA_FORMATS:
        _convert_data(input_path, output_path, src_ext, target)

    # ── Images ────────────────────────────────────────────────────────────────
    elif src_ext == 'heic':
        _convert_heic(input_path, output_path, target)
    elif src_ext == 'svg':
        _convert_svg(input_path, output_path, target)
    elif src_ext in RASTER_IMAGE_FORMATS:
        _convert_image(input_path, output_path, target)

    # ── Documents ─────────────────────────────────────────────────────────────
    elif src_ext == 'docx' and target == 'pdf':
        _convert_docx_to_pdf(input_path, output_path)
    elif src_ext == 'docx' and target == 'txt':
        _convert_docx_to_txt(input_path, output_path)
    elif src_ext == 'docx' and target == 'html':
        _convert_docx_to_html(input_path, output_path)
    elif src_ext == 'pdf' and target == 'docx':
        _convert_pdf_to_docx(input_path, output_path)
    elif src_ext == 'pdf' and target == 'txt':
        _convert_pdf_to_txt(input_path, output_path)
    elif src_ext == 'html' and target == 'pdf':
        _convert_html_to_pdf(input_path, output_path)
    elif src_ext == 'html' and target == 'md':
        _convert_html_to_md(input_path, output_path)
    elif src_ext == 'md' and target == 'html':
        _convert_md_to_html(input_path, output_path)

    # ── Video ─────────────────────────────────────────────────────────────────
    elif src_ext in VIDEO_FORMATS and target == 'gif':
        _convert_video_to_gif(input_path, output_path)
    elif src_ext in VIDEO_FORMATS and target == 'mp3':
        _extract_audio(input_path, output_path)
    elif src_ext in VIDEO_FORMATS and target in VIDEO_FORMATS:
        _convert_video(input_path, output_path)

    # ── Audio ─────────────────────────────────────────────────────────────────
    elif src_ext in AUDIO_FORMATS:
        _convert_audio(input_path, output_path, src_ext, target)

    # ── Archives ──────────────────────────────────────────────────────────────
    elif src_ext in ARCHIVE_FORMATS:
        _convert_archive(input_path, output_path, src_ext, target)


# ═══════════════════════════════════════════════════════════════════════════════
# DATA CONVERSIONS  (pandas)
# ═══════════════════════════════════════════════════════════════════════════════

def _convert_data(input_path: str, output_path: str, src: str, target: str) -> None:
    try:
        import pandas as pd
    except ImportError:
        raise ConversionError("pandas is not installed. Run: pip install pandas openpyxl lxml")

    try:
        if src == 'csv':
            df = pd.read_csv(input_path)
        elif src == 'json':
            df = pd.read_json(input_path)
        elif src == 'xml':
            try:
                df = pd.read_xml(input_path)
            except Exception:
                from lxml import etree
                root = etree.parse(input_path).getroot()
                rows = [{sub.tag: sub.text for sub in child} or {'value': child.text}
                        for child in root]
                df = pd.DataFrame(rows)
        elif src in ('xlsx', 'xls'):
            df = pd.read_excel(input_path)
        else:
            raise ConversionError(f"Unsupported source data format: {src}")
    except ConversionError:
        raise
    except Exception as exc:
        raise ConversionError(f"Failed to read {src.upper()} file: {exc}")

    try:
        if target == 'csv':
            df.to_csv(output_path, index=False)
        elif target == 'json':
            df.to_json(output_path, orient='records', indent=2, force_ascii=False)
        elif target == 'xlsx':
            df.to_excel(output_path, index=False, engine='openpyxl')
        else:
            raise ConversionError(f"Unsupported target data format: {target}")
    except ConversionError:
        raise
    except Exception as exc:
        raise ConversionError(f"Failed to write {target.upper()} file: {exc}")


# ═══════════════════════════════════════════════════════════════════════════════
# IMAGE CONVERSIONS
# ═══════════════════════════════════════════════════════════════════════════════

_PIL_FMT = {
    'jpg': 'JPEG', 'jpeg': 'JPEG',
    'png': 'PNG',  'gif': 'GIF',
    'bmp': 'BMP',  'webp': 'WEBP',
    'pdf': 'PDF',
}

def _convert_image(input_path: str, output_path: str, target: str) -> None:
    try:
        from PIL import Image
    except ImportError:
        raise ConversionError("Pillow is not installed. Run: pip install Pillow")

    pil_target = _PIL_FMT.get(target, target.upper())

    try:
        img = Image.open(input_path)

        if pil_target == 'PDF':
            # PDF requires RGB
            if img.mode not in ('RGB', 'L'):
                img = img.convert('RGB')
            img.save(output_path, 'PDF', resolution=150)
            return

        if pil_target == 'JPEG' and img.mode in ('RGBA', 'LA', 'P'):
            background = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'P':
                img = img.convert('RGBA')
            background.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
            img = background
        elif pil_target not in ('GIF',) and img.mode == 'P':
            img = img.convert('RGBA')

        save_kwargs = {}
        if pil_target == 'JPEG':
            save_kwargs['quality'] = 92
            if img.mode != 'RGB':
                img = img.convert('RGB')
        elif pil_target == 'WEBP':
            save_kwargs['quality'] = 90

        img.save(output_path, pil_target, **save_kwargs)
    except ConversionError:
        raise
    except Exception as exc:
        raise ConversionError(f"Image conversion failed: {exc}")


def _convert_heic(input_path: str, output_path: str, target: str) -> None:
    try:
        from pillow_heif import register_heif_opener
        register_heif_opener()
    except ImportError:
        raise ConversionError("pillow-heif is not installed. Run: pip install pillow-heif")
    # Once the opener is registered, delegate to the standard image converter
    _convert_image(input_path, output_path, target)


def _convert_svg(input_path: str, output_path: str, target: str) -> None:
    try:
        from svglib.svglib import svg2rlg
        from reportlab.graphics import renderPM, renderPDF
    except ImportError:
        raise ConversionError("svglib is not installed. Run: pip install svglib")

    try:
        drawing = svg2rlg(input_path)
        if drawing is None:
            raise ConversionError("Failed to parse SVG file — it may be malformed.")
        if target == 'png':
            renderPM.drawToFile(drawing, output_path, fmt='PNG')
        elif target == 'pdf':
            renderPDF.drawToFile(drawing, output_path)
    except ConversionError:
        raise
    except Exception as exc:
        raise ConversionError(f"SVG conversion failed: {exc}")


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT CONVERSIONS
# ═══════════════════════════════════════════════════════════════════════════════

# ── DOCX → PDF ────────────────────────────────────────────────────────────────

def _convert_docx_to_pdf(input_path: str, output_path: str) -> None:
    # Strategy 1: docx2pdf (MS Word / LibreOffice COM)
    try:
        from docx2pdf import convert
        convert(input_path, output_path)
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            return
    except ImportError:
        pass
    except Exception:
        pass

    # Strategy 2: python-docx + reportlab (text-only fallback)
    try:
        from docx import Document
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import cm

        doc    = Document(input_path)
        pdf    = SimpleDocTemplate(output_path, pagesize=A4,
                                   leftMargin=2*cm, rightMargin=2*cm,
                                   topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        story  = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                story.append(Spacer(1, 12))
                continue
            name = para.style.name or ''
            if   'Heading 1' in name: story.append(Paragraph(text, styles['h1']))
            elif 'Heading 2' in name: story.append(Paragraph(text, styles['h2']))
            elif 'Heading 3' in name: story.append(Paragraph(text, styles['h3']))
            else:                     story.append(Paragraph(text, styles['Normal']))

        pdf.build(story)
    except ImportError as exc:
        raise ConversionError(
            f"DOCX→PDF requires docx2pdf (+ MS Word/LibreOffice) or python-docx + reportlab.\n"
            f"Run: pip install docx2pdf  or  pip install python-docx reportlab\nDetail: {exc}"
        )
    except Exception as exc:
        raise ConversionError(f"DOCX→PDF failed: {exc}")


# ── DOCX → TXT ────────────────────────────────────────────────────────────────

def _convert_docx_to_txt(input_path: str, output_path: str) -> None:
    try:
        from docx import Document
    except ImportError:
        raise ConversionError("python-docx is not installed. Run: pip install python-docx")
    try:
        doc  = Document(input_path)
        text = '\n'.join(p.text for p in doc.paragraphs)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
    except Exception as exc:
        raise ConversionError(f"DOCX→TXT failed: {exc}")


# ── DOCX → HTML ───────────────────────────────────────────────────────────────

def _convert_docx_to_html(input_path: str, output_path: str) -> None:
    try:
        import mammoth
    except ImportError:
        raise ConversionError("mammoth is not installed. Run: pip install mammoth")
    try:
        with open(input_path, 'rb') as f:
            result = mammoth.convert_to_html(f)
        html = f'<!DOCTYPE html><html><head><meta charset="utf-8"></head><body>{result.value}</body></html>'
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html)
    except Exception as exc:
        raise ConversionError(f"DOCX→HTML failed: {exc}")


# ── PDF → DOCX ────────────────────────────────────────────────────────────────

def _convert_pdf_to_docx(input_path: str, output_path: str) -> None:
    try:
        from pdf2docx import Converter
    except ImportError:
        raise ConversionError("pdf2docx is not installed. Run: pip install pdf2docx")
    try:
        cv = Converter(input_path)
        cv.convert(output_path)
        cv.close()
    except Exception as exc:
        raise ConversionError(f"PDF→DOCX failed: {exc}")


# ── PDF → TXT ─────────────────────────────────────────────────────────────────

def _convert_pdf_to_txt(input_path: str, output_path: str) -> None:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ConversionError("PyMuPDF is not installed. Run: pip install pymupdf")
    try:
        doc   = fitz.open(input_path)
        pages = [page.get_text() for page in doc]
        doc.close()
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n\n'.join(pages))
    except Exception as exc:
        raise ConversionError(f"PDF→TXT failed: {exc}")


# ── HTML → PDF ────────────────────────────────────────────────────────────────

def _convert_html_to_pdf(input_path: str, output_path: str) -> None:
    # Strategy 1: weasyprint
    try:
        import weasyprint
        weasyprint.HTML(filename=input_path).write_pdf(output_path)
        return
    except ImportError:
        pass
    except Exception:
        pass

    # Strategy 2: xhtml2pdf
    try:
        from xhtml2pdf import pisa
        with open(input_path, 'rb') as src, open(output_path, 'wb') as dst:
            result = pisa.CreatePDF(src, dest=dst)
        if result.err:
            raise ConversionError("HTML→PDF: xhtml2pdf reported errors during conversion.")
    except ImportError:
        raise ConversionError(
            "HTML→PDF requires weasyprint or xhtml2pdf.\n"
            "Run: pip install weasyprint  or  pip install xhtml2pdf"
        )
    except ConversionError:
        raise
    except Exception as exc:
        raise ConversionError(f"HTML→PDF failed: {exc}")


# ── HTML → Markdown ───────────────────────────────────────────────────────────

def _convert_html_to_md(input_path: str, output_path: str) -> None:
    try:
        import markdownify
    except ImportError:
        raise ConversionError("markdownify is not installed. Run: pip install markdownify")
    try:
        with open(input_path, 'r', encoding='utf-8', errors='replace') as f:
            html = f.read()
        md = markdownify.markdownify(html, heading_style='ATX')
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md)
    except Exception as exc:
        raise ConversionError(f"HTML→Markdown failed: {exc}")


# ── Markdown → HTML ───────────────────────────────────────────────────────────

def _convert_md_to_html(input_path: str, output_path: str) -> None:
    try:
        import markdown
    except ImportError:
        raise ConversionError("markdown is not installed. Run: pip install markdown")
    try:
        with open(input_path, 'r', encoding='utf-8', errors='replace') as f:
            md_text = f.read()
        body = markdown.markdown(md_text, extensions=['tables', 'fenced_code'])
        html = (
            '<!DOCTYPE html><html><head>'
            '<meta charset="utf-8">'
            '<style>body{font-family:sans-serif;max-width:800px;margin:auto;padding:2rem}</style>'
            f'</head><body>{body}</body></html>'
        )
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html)
    except Exception as exc:
        raise ConversionError(f"Markdown→HTML failed: {exc}")


# ═══════════════════════════════════════════════════════════════════════════════
# VIDEO CONVERSIONS
# ═══════════════════════════════════════════════════════════════════════════════

def _convert_video_to_gif(input_path: str, output_path: str) -> None:
    try:
        from moviepy.editor import VideoFileClip
    except ImportError:
        raise ConversionError("moviepy is not installed. Run: pip install moviepy\nffmpeg must be on PATH.")

    try:
        clip = VideoFileClip(input_path)
        if clip.w > 480:
            clip = clip.resize(width=480)
        clip.write_gif(output_path, fps=min(clip.fps, 15), logger=None)
        clip.close()
    except ConversionError:
        raise
    except Exception as exc:
        raise ConversionError(f"Video→GIF failed: {exc}")


def _extract_audio(input_path: str, output_path: str) -> None:
    try:
        from moviepy.editor import VideoFileClip
    except ImportError:
        raise ConversionError("moviepy is not installed. Run: pip install moviepy\nffmpeg must be on PATH.")

    try:
        clip = VideoFileClip(input_path)
        if clip.audio is None:
            raise ConversionError("This video file has no audio track.")
        clip.audio.write_audiofile(output_path, logger=None)
        clip.close()
    except ConversionError:
        raise
    except Exception as exc:
        raise ConversionError(f"Audio extraction failed: {exc}")


def _convert_video(input_path: str, output_path: str) -> None:
    """Transcode between video container formats using ffmpeg directly."""
    try:
        result = subprocess.run(
            ['ffmpeg', '-i', input_path, '-y', output_path],
            capture_output=True, text=True, timeout=300,
        )
        if result.returncode != 0:
            raise ConversionError(f"ffmpeg error: {result.stderr[-400:]}")
    except FileNotFoundError:
        raise ConversionError(
            "ffmpeg is not on your PATH. Download from https://ffmpeg.org/download.html"
        )
    except subprocess.TimeoutExpired:
        raise ConversionError("Video conversion timed out (>5 min).")
    except ConversionError:
        raise
    except Exception as exc:
        raise ConversionError(f"Video conversion failed: {exc}")


# ═══════════════════════════════════════════════════════════════════════════════
# AUDIO CONVERSIONS  (pydub)
# ═══════════════════════════════════════════════════════════════════════════════

_PYDUB_FMT = {
    'mp3': 'mp3', 'wav': 'wav', 'ogg': 'ogg',
    'flac': 'flac', 'aac': 'adts', 'm4a': 'mp4',
}

def _convert_audio(input_path: str, output_path: str, src: str, target: str) -> None:
    try:
        from pydub import AudioSegment
    except ImportError:
        raise ConversionError("pydub is not installed. Run: pip install pydub\nffmpeg must be on PATH.")

    try:
        audio = AudioSegment.from_file(input_path, format=_PYDUB_FMT.get(src, src))
        audio.export(output_path, format=_PYDUB_FMT.get(target, target))
    except ConversionError:
        raise
    except Exception as exc:
        raise ConversionError(f"Audio conversion failed: {exc}")


# ═══════════════════════════════════════════════════════════════════════════════
# ARCHIVE CONVERSIONS  (stdlib zipfile + tarfile, rarfile for RAR)
# ═══════════════════════════════════════════════════════════════════════════════

def _convert_archive(input_path: str, output_path: str, src: str, target: str) -> None:
    if   src == 'zip' and target == 'tar': _zip_to_tar(input_path, output_path)
    elif src == 'tar' and target == 'zip': _tar_to_zip(input_path, output_path)
    elif src == 'rar' and target == 'zip': _rar_to_zip(input_path, output_path)


def _zip_to_tar(input_path: str, output_path: str) -> None:
    import zipfile, tarfile
    try:
        with zipfile.ZipFile(input_path, 'r') as zf:
            with tarfile.open(output_path, 'w') as tf:
                for name in zf.namelist():
                    data = zf.read(name)
                    info = tarfile.TarInfo(name=name)
                    info.size = len(data)
                    tf.addfile(info, io.BytesIO(data))
    except Exception as exc:
        raise ConversionError(f"ZIP→TAR failed: {exc}")


def _tar_to_zip(input_path: str, output_path: str) -> None:
    import zipfile, tarfile
    try:
        with tarfile.open(input_path, 'r:*') as tf:
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zf:
                for member in tf.getmembers():
                    if member.isfile():
                        f = tf.extractfile(member)
                        if f:
                            zf.writestr(member.name, f.read())
    except Exception as exc:
        raise ConversionError(f"TAR→ZIP failed: {exc}")


def _rar_to_zip(input_path: str, output_path: str) -> None:
    try:
        import rarfile
    except ImportError:
        raise ConversionError(
            "rarfile is not installed. Run: pip install rarfile\n"
            "Also needs unrar on PATH — on Windows 10/11 try: rarfile.UNRAR_TOOL = 'bsdtar'"
        )
    import zipfile
    try:
        with rarfile.RarFile(input_path, 'r') as rf:
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zf:
                for name in rf.namelist():
                    zf.writestr(name, rf.read(name))
    except Exception as exc:
        raise ConversionError(f"RAR→ZIP failed: {exc}")
